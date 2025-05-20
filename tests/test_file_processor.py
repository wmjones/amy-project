"""
Unit tests for file processing engine.
"""

import pytest
import tempfile
from pathlib import Path
from PIL import Image
import PyPDF2
from docx import Document
from src.file_access.processor import FileProcessor, ProcessedFile
import os


class TestFileProcessor:
    """Test FileProcessor functionality."""

    @pytest.fixture
    def processor(self):
        """Create a file processor instance."""
        return FileProcessor(
            max_chunk_size=1000, enable_ocr=False  # Disable OCR for unit tests
        )

    @pytest.fixture
    def sample_files(self):
        """Create sample test files."""
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)

            # Create text file
            text_file = temp_path / "sample.txt"
            text_file.write_text("This is a test document.\nWith multiple lines.\n")

            # Create image file
            image_file = temp_path / "sample.jpg"
            image = Image.new("RGB", (100, 100), color="red")
            image.save(image_file, "JPEG")

            # Create simple PDF
            pdf_file = temp_path / "sample.pdf"
            pdf_writer = PyPDF2.PdfWriter()
            page = pdf_writer.add_blank_page(width=200, height=200)
            pdf_writer.write(open(pdf_file, "wb"))

            # Create DOCX file
            docx_file = temp_path / "sample.docx"
            doc = Document()
            doc.add_paragraph("Test document content")
            doc.add_paragraph("Second paragraph")
            doc.save(docx_file)

            yield {
                "text": text_file,
                "image": image_file,
                "pdf": pdf_file,
                "docx": docx_file,
                "temp_dir": temp_path,
            }

    def test_process_text_file(self, processor, sample_files):
        """Test processing text files."""
        result = processor.process_file(str(sample_files["text"]))

        assert isinstance(result, ProcessedFile)
        assert result.success is True
        assert result.format == "text"
        assert "test document" in result.content.lower()
        assert result.metadata["lines"] == 3

    def test_process_image_file(self, processor, sample_files):
        """Test processing image files."""
        result = processor.process_file(str(sample_files["image"]))

        assert result.success is True
        assert result.format == "image"
        assert result.images is not None
        assert len(result.images) == 1
        assert result.metadata["format"] == "JPEG"
        assert result.metadata["size"] == (100, 100)

    def test_process_pdf_file(self, processor, sample_files):
        """Test processing PDF files."""
        # Create a PDF with text
        pdf_file = sample_files["temp_dir"] / "text.pdf"
        pdf_writer = PyPDF2.PdfWriter()
        # Note: PdfWriter in PyPDF2 doesn't support adding text directly in tests
        # This will create an empty PDF
        pdf_writer.add_blank_page(width=200, height=200)
        with open(pdf_file, "wb") as f:
            pdf_writer.write(f)

        result = processor.process_file(str(pdf_file))

        assert result.success is True
        assert result.format == "pdf"
        assert result.metadata["pages"] == 1

    def test_process_docx_file(self, processor, sample_files):
        """Test processing DOCX files."""
        result = processor.process_file(str(sample_files["docx"]))

        assert result.success is True
        assert result.format == "docx"
        assert "Test document content" in result.content
        assert "Second paragraph" in result.content
        assert result.metadata["paragraphs"] == 2

    def test_process_unsupported_file(self, processor, sample_files):
        """Test processing unsupported file types."""
        unsupported_file = sample_files["temp_dir"] / "test.xyz"
        unsupported_file.write_text("content")

        result = processor.process_file(str(unsupported_file))

        assert result.success is False
        assert "Unsupported file type" in result.error

    def test_process_nonexistent_file(self, processor):
        """Test processing non-existent files."""
        result = processor.process_file("/path/to/nonexistent/file.txt")

        assert result.success is False
        assert "File not found" in result.error

    def test_text_chunking(self, processor, sample_files):
        """Test text chunking for large files."""
        # Create large text file
        large_file = sample_files["temp_dir"] / "large.txt"
        large_content = "This is a test paragraph.\n\n" * 100
        large_file.write_text(large_content)

        result = processor.process_file(str(large_file))

        assert result.success is True
        assert result.chunks is not None
        assert len(result.chunks) > 1
        assert result.metadata["chunked"] is True

    def test_batch_processing(self, processor, sample_files):
        """Test batch file processing."""
        files = [
            str(sample_files["text"]),
            str(sample_files["image"]),
            str(sample_files["docx"]),
        ]

        results = processor.batch_process(files)

        assert len(results) == 3
        assert all(isinstance(r, ProcessedFile) for r in results)
        assert results[0].format == "text"
        assert results[1].format == "image"
        assert results[2].format == "docx"

    def test_encoding_detection(self, processor, sample_files):
        """Test text file encoding detection."""
        # Create file with specific encoding
        encoded_file = sample_files["temp_dir"] / "encoded.txt"
        content = "Héllo Wörld"
        encoded_file.write_text(content, encoding="utf-8")

        result = processor.process_file(str(encoded_file))

        assert result.success is True
        assert result.content == content
        assert result.metadata["encoding"] == "utf-8"

    def test_image_resizing(self, sample_files):
        """Test image resizing for large images."""
        # Create processor with specific max width
        processor = FileProcessor(image_max_width=50, enable_ocr=False)

        # Create large image
        large_image = sample_files["temp_dir"] / "large.jpg"
        image = Image.new("RGB", (200, 200), color="blue")
        image.save(large_image, "JPEG")

        result = processor.process_file(str(large_image))

        assert result.success is True
        assert result.metadata.get("resized") is True
        assert result.metadata["new_size"][0] == 50  # Width should be 50

    def test_error_handling(self, processor, sample_files):
        """Test error handling in file processing."""
        # Create corrupted file
        corrupt_file = sample_files["temp_dir"] / "corrupt.pdf"
        corrupt_file.write_text("This is not a valid PDF")

        result = processor.process_file(str(corrupt_file))

        assert result.success is False
        assert result.error is not None

    def test_progress_callback(self, processor, sample_files):
        """Test progress callback in batch processing."""
        files = [str(sample_files["text"]), str(sample_files["image"])]
        progress_updates = []

        def progress_callback(current, total):
            progress_updates.append((current, total))

        processor.batch_process(files, progress_callback=progress_callback)

        assert len(progress_updates) == 2
        assert progress_updates[0] == (1, 2)
        assert progress_updates[1] == (2, 2)
