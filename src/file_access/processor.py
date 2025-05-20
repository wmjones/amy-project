"""
File processing engine for various file formats.
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from dataclasses import dataclass
import io

# File format handlers
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import PyPDF2
from pdf2image import convert_from_path
from docx import Document
import chardet

# Base64 encoding for images
import base64

logger = logging.getLogger(__name__)


@dataclass
class ProcessedFile:
    """Container for processed file data."""

    content: str
    format: str
    metadata: Dict[str, Any]
    success: bool
    error: Optional[str] = None
    chunks: Optional[List[str]] = None
    images: Optional[List[str]] = None  # Base64 encoded images


class FileProcessor:
    """Main file processing engine."""

    def __init__(
        self,
        max_chunk_size: int = 10000,
        image_max_width: int = 1920,
        image_quality: int = 85,
        enable_ocr: bool = True,
    ):
        """Initialize file processor.

        Args:
            max_chunk_size: Maximum size for text chunks
            image_max_width: Maximum width for processed images
            image_quality: JPEG quality for compressed images
            enable_ocr: Whether to enable OCR for images
        """
        self.max_chunk_size = max_chunk_size
        self.image_max_width = image_max_width
        self.image_quality = image_quality
        self.enable_ocr = enable_ocr

        # Map extensions to processing methods
        self.processors = {
            ".jpg": self._process_image,
            ".jpeg": self._process_image,
            ".png": self._process_image,
            ".tiff": self._process_image,
            ".tif": self._process_image,
            ".bmp": self._process_image,
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".doc": self._process_docx,  # Will try with python-docx
            ".txt": self._process_text,
            ".text": self._process_text,
        }

        logger.info("File processor initialized")

    def process_file(self, file_path: str) -> ProcessedFile:
        """Process a file based on its type.

        Args:
            file_path: Path to the file

        Returns:
            ProcessedFile object with extracted content
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return ProcessedFile(
                content="",
                format="unknown",
                metadata={"error": "File not found"},
                success=False,
                error=f"File not found: {file_path}",
            )

        extension = file_path.suffix.lower()
        processor = self.processors.get(extension)

        if not processor:
            return ProcessedFile(
                content="",
                format=extension,
                metadata={"error": "Unsupported file type"},
                success=False,
                error=f"Unsupported file type: {extension}",
            )

        try:
            logger.info(f"Processing {file_path} with {processor.__name__}")
            return processor(file_path)
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return ProcessedFile(
                content="",
                format=extension,
                metadata={"error": str(e)},
                success=False,
                error=str(e),
            )

    def _process_image(self, file_path: Path) -> ProcessedFile:
        """Process image files."""
        try:
            with Image.open(file_path) as img:
                # Get basic metadata
                metadata = {
                    "format": img.format,
                    "size": img.size,
                    "mode": img.mode,
                    "has_transparency": img.mode in ("RGBA", "LA"),
                }

                # Convert to RGB if necessary
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")

                # Resize if too large
                if img.width > self.image_max_width:
                    ratio = self.image_max_width / img.width
                    new_size = (int(img.width * ratio), int(img.height * ratio))
                    img = img.resize(new_size, Image.Resampling.LANCZOS)
                    metadata["resized"] = True
                    metadata["new_size"] = new_size

                # Enhance image for better OCR
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(1.5)

                # Convert to base64 for Claude
                buffered = io.BytesIO()
                img.save(buffered, format="JPEG", quality=self.image_quality)
                img_base64 = base64.b64encode(buffered.getvalue()).decode()

                # Perform OCR if enabled
                text_content = ""
                if self.enable_ocr:
                    try:
                        text_content = pytesseract.image_to_string(img)
                        metadata["ocr_performed"] = True
                        metadata["ocr_text_length"] = len(text_content)
                    except Exception as e:
                        logger.warning(f"OCR failed for {file_path}: {e}")
                        metadata["ocr_error"] = str(e)

                return ProcessedFile(
                    content=text_content,
                    format="image",
                    metadata=metadata,
                    success=True,
                    images=[img_base64],
                )

        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}")
            raise

    def _process_pdf(self, file_path: Path) -> ProcessedFile:
        """Process PDF files."""
        try:
            text_content = []
            metadata = {"pages": 0, "has_images": False, "is_scanned": False}

            # Try text extraction first
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)

                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{text}")

                # Check if PDF has extractable text
                if not any(text_content):
                    metadata["is_scanned"] = True

            # If no text found and OCR is enabled, try OCR on PDF pages
            if not text_content and self.enable_ocr:
                try:
                    logger.info(f"PDF appears to be scanned. Attempting OCR...")
                    images = convert_from_path(str(file_path), dpi=200)

                    for i, image in enumerate(images):
                        # Enhance image for OCR
                        enhancer = ImageEnhance.Contrast(image)
                        image = enhancer.enhance(1.5)

                        # Perform OCR
                        text = pytesseract.image_to_string(image)
                        if text.strip():
                            text_content.append(f"--- Page {i + 1} (OCR) ---\n{text}")

                    metadata["ocr_performed"] = True

                except Exception as e:
                    logger.warning(f"OCR on PDF failed: {e}")
                    metadata["ocr_error"] = str(e)

            # Join all text content
            full_text = "\n\n".join(text_content) if text_content else ""

            # Chunk text if needed
            chunks = None
            if len(full_text) > self.max_chunk_size:
                chunks = self._chunk_text(full_text)
                metadata["chunked"] = True
                metadata["chunk_count"] = len(chunks)

            return ProcessedFile(
                content=full_text,
                format="pdf",
                metadata=metadata,
                success=True,
                chunks=chunks,
            )

        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise

    def _process_docx(self, file_path: Path) -> ProcessedFile:
        """Process DOCX files."""
        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_content.append(cell.text.strip())
                    if row_content:
                        table_content.append(" | ".join(row_content))
                if table_content:
                    tables_text.append("\n".join(table_content))

            # Combine all text
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n--- Tables ---\n\n" + "\n\n".join(tables_text)

            # Metadata
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
            }

            # Try to extract properties
            try:
                core_props = doc.core_properties
                metadata.update(
                    {
                        "author": core_props.author,
                        "created": (
                            str(core_props.created) if core_props.created else None
                        ),
                        "modified": (
                            str(core_props.modified) if core_props.modified else None
                        ),
                        "title": core_props.title,
                    }
                )
            except Exception as e:
                logger.warning(f"Could not extract document properties: {e}")

            # Chunk if needed
            chunks = None
            if len(full_text) > self.max_chunk_size:
                chunks = self._chunk_text(full_text)
                metadata["chunked"] = True
                metadata["chunk_count"] = len(chunks)

            return ProcessedFile(
                content=full_text,
                format="docx",
                metadata=metadata,
                success=True,
                chunks=chunks,
            )

        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

    def _process_text(self, file_path: Path) -> ProcessedFile:
        """Process text files."""
        try:
            # Detect encoding
            with open(file_path, "rb") as file:
                raw_data = file.read()
                encoding_info = chardet.detect(raw_data)
                encoding = encoding_info["encoding"] or "utf-8"

            # Read text with detected encoding
            with open(file_path, "r", encoding=encoding) as file:
                content = file.read()

            # Metadata
            metadata = {
                "encoding": encoding,
                "confidence": encoding_info["confidence"],
                "size": len(content),
                "lines": content.count("\n") + 1,
            }

            # Chunk if needed
            chunks = None
            if len(content) > self.max_chunk_size:
                chunks = self._chunk_text(content)
                metadata["chunked"] = True
                metadata["chunk_count"] = len(chunks)

            return ProcessedFile(
                content=content,
                format="text",
                metadata=metadata,
                success=True,
                chunks=chunks,
            )

        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise

    def _chunk_text(self, text: str, overlap: int = 200) -> List[str]:
        """Split text into chunks with overlap.

        Args:
            text: Text to chunk
            overlap: Number of characters to overlap between chunks

        Returns:
            List of text chunks
        """
        chunks = []

        # Split by paragraphs first to maintain context
        paragraphs = text.split("\n\n")
        current_chunk = []
        current_size = 0

        for para in paragraphs:
            para_size = len(para)

            # If adding this paragraph exceeds limit, save current chunk
            if current_size + para_size > self.max_chunk_size and current_chunk:
                chunks.append("\n\n".join(current_chunk))

                # Keep last paragraph for overlap
                if overlap > 0 and current_chunk:
                    current_chunk = [current_chunk[-1]]
                    current_size = len(current_chunk[0])
                else:
                    current_chunk = []
                    current_size = 0

            current_chunk.append(para)
            current_size += para_size + 2  # +2 for \n\n

        # Add remaining content
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return chunks

    def batch_process(
        self, file_paths: List[str], progress_callback=None
    ) -> List[ProcessedFile]:
        """Process multiple files in batch.

        Args:
            file_paths: List of file paths to process
            progress_callback: Optional callback for progress updates

        Returns:
            List of ProcessedFile objects
        """
        results = []
        total_files = len(file_paths)

        for i, file_path in enumerate(file_paths):
            result = self.process_file(file_path)
            results.append(result)

            if progress_callback:
                progress_callback(i + 1, total_files)

        return results
